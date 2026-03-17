import html
import html
import re
from io import BytesIO
from typing import Any, Dict, List, Optional

import numpy as np
import pandas as pd
import streamlit as st
import streamlit.components.v1 as components
from docx import Document

from src.core.category_rules import norm_indicator_name, order_categories_like_ui
from src.core.formatters import fmt_num, fmt_period, fmt_triangle_delta
from src.features.insights import (
    build_activity_snapshot,
    compute_anomaly_table,
    compute_contribution_table,
    compute_gyeonggi_vs_national_contribution,
    fmt_contrib_items,
    infer_lag_from_df,
)
from src.features.new_history import collect_new_events

ALLOWED_INDUSTRY_FACTOR_CODES = {"A", "C", "F", "GI", "EL~U", "DHJK"}


def _industry_code_from_label(label: object) -> str:
    s = str(label or "").strip().upper()
    if not s:
        return ""

    raw_text = str(label or "").strip()
    norm_text = re.sub(r"[\s\u00A0·ㆍ,./()\-]", "", raw_text)
    if norm_text in {"\uACC4", "\uD569\uACC4", "\uC804\uCCB4"} or s == "TOTAL":
        return "TOTAL"

    def _normalize_token(raw: str) -> str:
        token_raw = re.sub(r"[^A-Z~,\-]", "", str(raw or "").upper())
        if not token_raw or not re.search(r"[A-Z]", token_raw):
            return ""
        normalized = token_raw.replace(",", "").replace("-", "~")
        letters_only = re.sub(r"[^A-Z]", "", normalized)
        if letters_only == "ELU" or normalized == "EL~U":
            return "EL~U"
        if letters_only == "DU" or normalized == "D~U":
            return "D~U"
        if letters_only in {"BC", "GI", "DHJK", "A", "C", "F"}:
            return letters_only
        return letters_only

    m = re.search(r"\(([^()]*)\)", s)
    if m:
        token = _normalize_token(m.group(1))
        if token:
            return token
    m = re.match(r"^[^A-Z0-9]*([A-Z])", s)
    if m:
        return m.group(1)
    return ""


def _filter_industry_factor_table(table: pd.DataFrame) -> pd.DataFrame:
    if table is None or table.empty or "분류" not in table.columns:
        return table.copy() if isinstance(table, pd.DataFrame) else pd.DataFrame()
    view = table.copy()
    view["_industry_code"] = view["분류"].map(_industry_code_from_label)
    view = view[view["_industry_code"].isin(ALLOWED_INDUSTRY_FACTOR_CODES)].copy()
    return view.drop(columns=["_industry_code"], errors="ignore")


def _current_streak_length(values: pd.Series, positive: bool) -> int:
    if values.empty:
        return 0
    cnt = 0
    for v in reversed(values.tolist()):
        if pd.isna(v):
            break
        vv = float(v)
        if positive and vv > 0:
            cnt += 1
            continue
        if (not positive) and vv < 0:
            cnt += 1
            continue
        break
    return cnt


def _collect_dataset_streak_items(
    source_df: pd.DataFrame,
    region: str,
    dataset_key: str,
    asof_period: pd.Timestamp,
    indicator_name: str = "",
    min_len: int = 3,
) -> tuple[List[Dict[str, object]], List[Dict[str, object]], str]:
    if source_df.empty or not region or pd.isna(asof_period):
        return [], [], "개월"
    ds = source_df[
        (source_df["dataset_key"].astype(str) == str(dataset_key))
        & (source_df["region_name"].astype(str) == str(region))
        & (pd.to_datetime(source_df["period"], errors="coerce") <= pd.Timestamp(asof_period))
    ].copy()
    if ds.empty or "yoy_abs" not in ds.columns:
        return [], [], "개월"
    ds["period"] = pd.to_datetime(ds["period"], errors="coerce")
    ds = ds.dropna(subset=["period"])
    if ds.empty:
        return [], [], "개월"
    if indicator_name:
        ds = ds[ds["indicator_name"].astype(str) == str(indicator_name)].copy()
    if ds.empty:
        return [], [], "개월"

    ds_prd_se = (
        str(ds["prd_se"].dropna().iloc[0]).upper()
        if "prd_se" in ds.columns and not ds["prd_se"].dropna().empty
        else "M"
    )
    dur_unit = "반기" if ds_prd_se == "H" else "개월"
    up_items: List[Dict[str, object]] = []
    down_items: List[Dict[str, object]] = []

    for _, g in ds.groupby(["indicator_name", "category_name"], dropna=False):
        g = g.sort_values("period")
        if g.empty:
            continue
        latest_period = pd.Timestamp(g["period"].iloc[-1])
        if latest_period != pd.Timestamp(asof_period):
            continue
        yoy_values = pd.to_numeric(g["yoy_abs"], errors="coerce")
        if yoy_values.empty or pd.isna(yoy_values.iloc[-1]):
            continue
        latest_yoy = float(yoy_values.iloc[-1])
        cat = str(g["category_name"].iloc[0]).strip()
        ind = str(g["indicator_name"].iloc[0]).strip()
        label = cat if cat else (ind if ind else "전체")

        up_len = _current_streak_length(yoy_values, positive=True)
        if up_len >= min_len:
            up_items.append(
                {
                    "label": label,
                    "len": int(up_len),
                    "start": pd.Timestamp(g["period"].iloc[len(g) - up_len]),
                    "end": latest_period,
                    "latest_yoy": latest_yoy,
                    "unit": dur_unit,
                    "dir": "증가",
                }
            )
        down_len = _current_streak_length(yoy_values, positive=False)
        if down_len >= min_len:
            down_items.append(
                {
                    "label": label,
                    "len": int(down_len),
                    "start": pd.Timestamp(g["period"].iloc[len(g) - down_len]),
                    "end": latest_period,
                    "latest_yoy": latest_yoy,
                    "unit": dur_unit,
                    "dir": "감소",
                }
            )
    return up_items, down_items, dur_unit


def _build_dataset_streak_summary_line(
    source_df: pd.DataFrame,
    region: str,
    dataset_key: str,
    asof_period: pd.Timestamp,
    indicator_name: str = "",
    min_len: int = 3,
    yoy_label: str = "전년동월",
) -> str:
    up_items, down_items, dur_unit = _collect_dataset_streak_items(
        source_df=source_df,
        region=region,
        dataset_key=dataset_key,
        asof_period=asof_period,
        indicator_name=indicator_name,
        min_len=min_len,
    )

    if not up_items and not down_items:
        return f"연속 증가/감소 요약(3{dur_unit} 이상): 없음 ({yoy_label}대비 증감 기준)"

    sort_key = lambda x: (-int(x["len"]), -abs(float(x["latest_yoy"])), str(x["label"]))
    up_items = sorted(up_items, key=sort_key)
    down_items = sorted(down_items, key=sort_key)
    ordered_items = up_items + down_items

    tokens: List[str] = []
    for item in ordered_items:
        s_txt = str(fmt_period(item["start"], "H" if item["unit"] == "반기" else "M"))
        e_txt = str(fmt_period(item["end"], "H" if item["unit"] == "반기" else "M"))
        l_txt = str(item["label"])
        tokens.append(f"{l_txt} {s_txt}~{e_txt} {item['len']}{item['unit']} 연속 {item['dir']}")
    return f"연속 증가/감소 요약(3{ordered_items[0]['unit']} 이상): " + ", ".join(tokens) + f" ({yoy_label}대비 증감 기준)"


def _to_docx_text(text: object) -> str:
    return str(text).replace("\\", "")


def _add_docx_table(doc: Document, df: pd.DataFrame) -> None:
    if df is None or df.empty:
        doc.add_paragraph("데이터 없음")
        return
    view = df.copy()
    table = doc.add_table(rows=1, cols=len(view.columns))
    table.style = "Table Grid"
    for i, col in enumerate(view.columns):
        table.rows[0].cells[i].text = _to_docx_text(col)
    for _, row in view.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(view.columns):
            v = row[col]
            cells[i].text = "-" if pd.isna(v) else _to_docx_text(v)


def _build_report_docx_bytes(
    title: str,
    summary_lines: List[str],
    activity_table: pd.DataFrame,
    structure_lines: List[str],
    reference_title: str,
    reference_lines: List[str],
    detail_sections: List[Dict[str, Any]],
) -> bytes:
    doc = Document()
    doc.add_heading(_to_docx_text(title), level=0)
    doc.add_paragraph("(출처:경제활동인구조사, 통계청)")

    doc.add_heading("##월간 핵심요약", level=1)
    for line in summary_lines:
        doc.add_paragraph(_to_docx_text(line), style="List Bullet")

    doc.add_heading("##경제활동인구 현황요약", level=1)
    _add_docx_table(doc, activity_table)

    doc.add_heading("##취업자수 상세현황", level=1)
    for line in structure_lines:
        doc.add_paragraph(_to_docx_text(line), style="List Bullet")

    doc.add_heading(_to_docx_text(reference_title), level=1)
    for line in reference_lines:
        doc.add_paragraph(_to_docx_text(line), style="List Bullet")

    doc.add_page_break()
    doc.add_heading("[참고] 취업자수 상세내용", level=0)
    for section in detail_sections:
        doc.add_heading(_to_docx_text(section.get("title", "")), level=1)
        for line in section.get("lines", []):
            doc.add_paragraph(_to_docx_text(line), style="List Bullet")
        _add_docx_table(doc, section.get("table", pd.DataFrame()))

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _clean_html_report_text(text: object) -> str:
    s = str(text)
    s = re.sub(r"\\([\\`*_{}\[\]()#+\-.!|~])", r"\1", s)
    s = s.replace("\\", "")
    return s


def _lines_to_html(lines: List[str]) -> str:
    if not lines:
        return "<p>??? ??</p>"
    items = "".join(f"<li>{html.escape(_clean_html_report_text(line))}</li>" for line in lines)
    return f"<ul>{items}</ul>"


def _table_to_html(df: pd.DataFrame) -> str:
    if df is None or df.empty:
        return "<p>??? ??</p>"
    view = df.copy().fillna("-")
    view.columns = [_clean_html_report_text(c) for c in view.columns]
    view = view.astype(str).applymap(_clean_html_report_text)
    return view.to_html(index=False, escape=True, classes="brief-table")


def _build_printable_report_html(
    title: str,
    summary_lines: List[str],
    activity_table: pd.DataFrame,
    structure_lines: List[str],
    reference_title: str,
    reference_lines: List[str],
    detail_sections: List[Dict[str, Any]],
) -> str:
    detail_html = []
    for section in detail_sections:
        sec_title = html.escape(_clean_html_report_text(section.get("title", "")))
        sec_lines = _lines_to_html([str(x) for x in section.get("lines", [])])
        sec_table = _table_to_html(section.get("table", pd.DataFrame()))
        detail_html.append(
            f"<section class='section-block detail-block'><h3>{sec_title}</h3>{sec_lines}{sec_table}</section>"
        )
    detail_joined = "".join(detail_html)

    return f"""
<!doctype html>
<html>
<head>
  <meta charset="utf-8" />
  <style>
    @import url('https://fonts.googleapis.com/css2?family=Nanum+Gothic:wght@400;700;800&display=swap');
    :root {{
      --navy-900: #0f2742;
      --navy-700: #1e3a5f;
      --navy-500: #2f567f;
      --line: #d6dde8;
      --text: #111827;
      --muted: #5f6f86;
      --paper: #ffffff;
      --soft: #f4f7fb;
    }}
    body {{
      font-family: 'Nanum Gothic', 'Malgun Gothic', 'Apple SD Gothic Neo', 'Noto Sans KR', sans-serif;
      color: var(--text);
      margin: 0;
      padding: 20px;
      line-height: 1.6;
      background: var(--soft);
    }}
    .sheet {{
      max-width: 980px;
      margin: 0 auto;
      background: var(--paper);
      border: 1px solid var(--line);
      border-radius: 14px;
      padding: 22px 24px;
      box-shadow: 0 12px 30px rgba(15, 39, 66, 0.08);
    }}
    .toolbar {{
      display: flex;
      gap: 10px;
      align-items: center;
      margin-bottom: 14px;
    }}
    .print-btn {{
      border: 1px solid var(--navy-700);
      color: #fff;
      background: linear-gradient(180deg, var(--navy-500), var(--navy-700));
      padding: 7px 12px;
      border-radius: 8px;
      cursor: pointer;
      font-size: 13px;
      font-weight: 700;
    }}
    .hint {{
      color: var(--muted);
      font-size: 12px;
    }}
    h1 {{
      margin: 0;
      font-size: 33px;
      color: var(--navy-900);
      text-decoration: underline;
      text-decoration-thickness: 2px;
      text-underline-offset: 3px;
    }}
    h2 {{
      margin: 18px 0 10px 0;
      color: var(--navy-700);
      font-size: 27px;
      border-left: 6px solid var(--navy-700);
      padding-left: 10px;
    }}
    h3 {{
      margin: 0 0 8px 0;
      font-size: 19px;
      color: var(--navy-900);
    }}
    ul {{
      margin: 4px 0 4px 18px;
      padding: 0;
    }}
    li {{
      margin: 4px 0;
    }}
    li::marker {{
      color: var(--navy-700);
      font-weight: 700;
    }}
    .section-block {{
      border: 1px solid var(--line);
      border-radius: 10px;
      padding: 12px 14px;
      margin: 10px 0 14px 0;
      background: #fff;
    }}
    .detail-block {{
      margin-top: 12px;
    }}
    table.brief-table {{
      border-collapse: collapse;
      width: 100%;
      margin: 8px 0 14px 0;
      table-layout: fixed;
      font-size: 13px;
    }}
    table.brief-table th, table.brief-table td {{
      border: 1px solid var(--line);
      padding: 7px 8px;
      text-align: center;
      word-break: break-word;
    }}
    table.brief-table th {{
      background: var(--navy-700);
      color: #fff;
      font-weight: 700;
    }}
    table.brief-table tr:nth-child(even) td {{
      background: #f8fbff;
    }}
    hr {{
      border: none;
      border-top: 1px solid var(--line);
      margin: 10px 0 16px 0;
    }}
    .source {{
      color: var(--muted);
      font-size: 13px;
      margin-top: 4px;
      margin-bottom: 8px;
    }}
    .page-break {{
      page-break-before: always;
    }}
    @media print {{
      .toolbar {{ display: none; }}
      body {{
        padding: 0;
        background: #fff;
      }}
      .sheet {{
        max-width: none;
        border: none;
        border-radius: 0;
        box-shadow: none;
        padding: 0;
      }}
      table.brief-table {{ font-size: 11px; }}
      @page {{ size: A4; margin: 12mm; }}
    }}
  </style>
</head>
<body>
  <article class="sheet">
    <div class="toolbar">
      <button class="print-btn" onclick="window.print()">인쇄 / PDF 저장</button>
      <span class="hint">브라우저 인쇄(Ctrl+P)에서 PDF 저장을 선택하세요.</span>
    </div>
    <h1>{html.escape(_clean_html_report_text(title))}</h1>
    <div class="source">(출처: 경제활동인구조사, 통계청)</div>
    <hr />
    <h2>##월간 핵심요약</h2>
    <section class="section-block">{_lines_to_html(summary_lines)}</section>
    <h2>##경제활동인구 현황요약</h2>
    <section class="section-block">{_table_to_html(activity_table)}</section>
    <h2>##취업자수 상세현황</h2>
    <section class="section-block">{_lines_to_html(structure_lines)}</section>
    <h2>{html.escape(_clean_html_report_text(reference_title))}</h2>
    <section class="section-block">{_lines_to_html(reference_lines)}</section>
    <div class="page-break"></div>
    <h2>[참고] 취업자수 상세내용</h2>
    {detail_joined}
  </article>
</body>
</html>
"""


def render_report_template(
    df: pd.DataFrame,
    province_df: pd.DataFrame,
    region_pool: List[str],
    datasets: List[Any],
    is_gyeonggi31_mode: bool,
    labels: Dict[str, str],
    for_pdf: bool = False,
) -> None:
    if df.empty:
        st.info("리포트 생성 대상 데이터가 없습니다.")
        return

    work_df = df.copy()
    work_df["period"] = pd.to_datetime(work_df["period"], errors="coerce")
    work_df = work_df.dropna(subset=["period"])
    if work_df.empty:
        st.info("리포트 생성 대상 기간 데이터가 없습니다.")
        return

    province_work_df = province_df.copy()
    if not province_work_df.empty:
        province_work_df["period"] = pd.to_datetime(province_work_df["period"], errors="coerce")
        province_work_df = province_work_df.dropna(subset=["period"])

    default_region = "경기도" if "경기도" in region_pool else (region_pool[0] if region_pool else "")
    if not default_region:
        st.info("리포트 생성 대상 지역이 없습니다.")
        return

    c_region, c_period = st.columns([1, 1])
    with c_region:
        region = st.selectbox(
            "리포트 지역",
            region_pool,
            index=region_pool.index(default_region) if default_region in region_pool else 0,
            key="report_template_region",
        )

    activity_base = work_df[(work_df["dataset_key"] == "activity") & (work_df["region_name"] == region)].copy()
    if activity_base.empty:
        st.info("선택한 지역의 경제활동인구현황 데이터가 없습니다.")
        return
    prd_se_for_period = (
        str(activity_base["prd_se"].dropna().iloc[0]).upper()
        if "prd_se" in activity_base.columns and not activity_base["prd_se"].dropna().empty
        else "M"
    )
    period_values = sorted(activity_base["period"].dropna().unique().tolist(), reverse=True)
    if not period_values:
        st.info("선택 가능한 기준시점이 없습니다.")
        return
    period_labels = [fmt_period(p, prd_se_for_period) for p in period_values]
    period_map = {label: pd.Timestamp(value) for label, value in zip(period_labels, period_values)}
    with c_period:
        selected_period_label = st.selectbox("기준시점", period_labels, index=0, key="report_template_period")
    selected_period = period_map[selected_period_label]

    report_df = work_df[work_df["period"] <= selected_period].copy()
    if report_df.empty:
        st.info("선택한 기준시점 이전 데이터가 없습니다.")
        return
    province_report_df = (
        province_work_df[province_work_df["period"] <= selected_period].copy()
        if not province_work_df.empty
        else pd.DataFrame()
    )

    lag = infer_lag_from_df(report_df)
    activity_df, activity_meta = build_activity_snapshot(report_df, region, lag)
    if not activity_meta.get("ok"):
        st.info(str(activity_meta.get("message", "경제활동인구 요약 데이터를 생성할 수 없습니다.")))
        return

    prd_se = str(activity_meta.get("prd_se", "M"))
    latest_period = activity_meta.get("latest_period")
    prev_period = activity_meta.get("prev_period")
    latest_text = fmt_period(latest_period, prd_se)
    prev_text = fmt_period(prev_period, prd_se)
    yoy_text = labels.get("yoy", "전년동월")
    st.caption(f"리포트 기준: {region} / {latest_text}")
    report_title = f"{latest_text} {region} 경제활동인구 브리프"
    st.markdown(
        f"<h2 style='margin-bottom:0.2rem;'><u>{report_title}</u> "
        f"<span style='font-size:0.55em; font-weight:500;'>(출처:경제활동인구조사, 통계청)</span></h2>",
        unsafe_allow_html=True,
    )
    st.markdown("<hr style='margin-top:0.1rem; margin-bottom:0.8rem;'>", unsafe_allow_html=True)

    def _report_heading(text: str) -> None:
        st.markdown(
            f"<p style='color:#2f6fb3; font-weight:800; font-size:1.9rem; margin:0.2rem 0;'>##{text}</p>",
            unsafe_allow_html=True,
        )

    def _get_row(norm_name: str) -> Optional[pd.Series]:
        view = activity_df[activity_df["norm_indicator"] == norm_name]
        if view.empty:
            return None
        return view.iloc[0]

    emp_row = _get_row(norm_indicator_name("취업자"))
    emp_rate_row = _get_row(norm_indicator_name("고용률"))
    unemp_rate_row = _get_row(norm_indicator_name("실업률"))

    industry_df, industry_meta = compute_contribution_table(report_df, region=region, dataset_key="industry", lag=lag)
    industry_factor_df = _filter_industry_factor_table(industry_df)
    occupation_df, occupation_meta = compute_contribution_table(report_df, region=region, dataset_key="occupation", lag=lag)

    emp_delta = float(emp_row["delta_value"]) if emp_row is not None and pd.notna(emp_row.get("delta_value")) else np.nan
    if pd.isna(emp_delta):
        emp_direction = "방향을 확인할 수 없습니다."
    elif emp_delta > 0:
        emp_direction = "증가했습니다."
    elif emp_delta < 0:
        emp_direction = "감소했습니다."
    else:
        emp_direction = "보합입니다."

    compare_base = province_report_df if not province_report_df.empty else report_df
    gy_meta_summary: Dict[str, Any] = {}
    if region != "전국":
        _, gy_meta_summary = compute_gyeonggi_vs_national_contribution(compare_base, region_name=region)

    def _top_factor_text(table: pd.DataFrame, unit: str, positive: bool) -> str:
        if table.empty:
            return "정보 없음"
        view = table[table["증감"] > 0] if positive else table[table["증감"] < 0]
        if view.empty:
            return "없음"
        row = view.nlargest(1, "증감").iloc[0] if positive else view.nsmallest(1, "증감").iloc[0]
        pct_text = "-" if pd.isna(row.get("기여율(%)")) else f"{float(row['기여율(%)']):,.1f}%"
        return f"{row['분류']}({fmt_num(row['증감'], unit)}, {pct_text})"

    report_events = collect_new_events(report_df)
    if not report_events.empty:
        report_events = report_events[
            (report_events["지역"].astype(str) == str(region))
            & (report_events["기준월"].astype(str) == str(latest_text))
        ].copy()

    def _fmt_streak_item(item: Dict[str, object]) -> str:
        prd_se_local = "H" if str(item.get("unit", "")) == "반기" else "M"
        s_txt = fmt_period(item.get("start"), prd_se_local)
        e_txt = fmt_period(item.get("end"), prd_se_local)
        return (
            f"{item.get('dataset', '')} {item.get('label', '전체')} "
            f"{s_txt}~{e_txt} {int(item.get('len', 0))}{item.get('unit', '')} 연속 {item.get('dir', '')}"
        )

    def _pick_top_streak_items() -> tuple[Optional[Dict[str, object]], Optional[Dict[str, object]]]:
        source_defs = [
            (
                "경제활동인구현황(취업자)",
                "activity",
                str(emp_row.get("지표", "취업자")) if emp_row is not None else "취업자",
            ),
            ("산업별 취업자수", "industry", str(industry_meta.get("indicator", ""))),
            ("직종별 취업자수", "occupation", str(occupation_meta.get("indicator", ""))),
            ("종사상지위별 취업자", "status", ""),
        ]
        up_pool: List[Dict[str, object]] = []
        down_pool: List[Dict[str, object]] = []
        for ds_title, ds_key, ind_name in source_defs:
            up_items, down_items, _ = _collect_dataset_streak_items(
                source_df=report_df,
                region=region,
                dataset_key=ds_key,
                asof_period=pd.Timestamp(selected_period),
                indicator_name=ind_name,
                min_len=3,
            )
            for item in up_items:
                tagged = dict(item)
                tagged["dataset"] = ds_title
                up_pool.append(tagged)
            for item in down_items:
                tagged = dict(item)
                tagged["dataset"] = ds_title
                down_pool.append(tagged)

        sort_key = lambda x: (
            -int(x.get("len", 0)),
            -abs(float(x.get("latest_yoy", 0.0))),
            str(x.get("dataset", "")),
            str(x.get("label", "")),
        )
        up_pool = sorted(up_pool, key=sort_key)
        down_pool = sorted(down_pool, key=sort_key)
        return (up_pool[0] if up_pool else None, down_pool[0] if down_pool else None)

    def _build_new_focus_line(event_type: str, line_title: str) -> str:
        if report_events.empty:
            return f"{line_title}: 없음"
        view = report_events[report_events["유형"].astype(str) == str(event_type)].copy()
        if view.empty:
            return f"{line_title}: 없음"

        ds_order = {
            "경제활동인구현황": 0,
            "산업별 취업자수": 1,
            "직종별 취업자수": 2,
            "종사상지위별 취업자": 3,
            "연령별 취업자": 4,
        }
        metric_order = {"원자료": 0, "YoY(절대)": 1, "YoY(증감률)": 2}
        scope_order = {"전체기간": 0, "최근5년": 1}
        view["_ds"] = view["데이터셋"].map(ds_order).fillna(99)
        view["_metric"] = view["구분"].map(metric_order).fillna(99)
        view["_scope"] = view["범위"].map(scope_order).fillna(99)
        view = view.sort_values(["_ds", "_metric", "_scope", "지표", "분류"])

        items: List[str] = []
        for _, row in view.head(2).iterrows():
            cat = str(row.get("분류", "")).strip()
            cat_text = cat if cat else "전체"
            items.append(
                f"{row['데이터셋']} {row['지표']}/{cat_text} ({row['구분']} {row['범위']} {row['유형']})"
            )
        return f"{line_title}: " + (", ".join(items) if items else "없음")

    _report_heading("월간 핵심요약")
    summary_lines = [
        f"{latest_text} {region} 취업자는 {fmt_num(emp_row['latest_value'] if emp_row is not None else np.nan, str(emp_row['unit']) if emp_row is not None else '')}로, "
        f"{prev_text} 대비 {fmt_num(emp_row['delta_value'] if emp_row is not None else np.nan, str(emp_row['unit']) if emp_row is not None else '')} {emp_direction}",
        f"고용률은 {fmt_num(emp_rate_row['latest_value'] if emp_rate_row is not None else np.nan, str(emp_rate_row['unit']) if emp_rate_row is not None else '%')}"
        f"({yoy_text} 대비 {fmt_num(emp_rate_row['delta_value'] if emp_rate_row is not None else np.nan, str(emp_rate_row['unit']) if emp_rate_row is not None else '%')}), "
        f"실업률은 {fmt_num(unemp_rate_row['latest_value'] if unemp_rate_row is not None else np.nan, str(unemp_rate_row['unit']) if unemp_rate_row is not None else '%')}"
        f"({yoy_text} 대비 {fmt_num(unemp_rate_row['delta_value'] if unemp_rate_row is not None else np.nan, str(unemp_rate_row['unit']) if unemp_rate_row is not None else '%')})입니다.",
        f"(산업) 증가요인 1위는 {_top_factor_text(industry_factor_df, str(industry_meta.get('unit', '')), True)}, "
        f"감소요인 1위는 {_top_factor_text(industry_factor_df, str(industry_meta.get('unit', '')), False)}입니다.",
        f"(직종) 증가요인 1위는 {_top_factor_text(occupation_df, str(industry_meta.get('unit', '')), True)}, "
        f"감소요인 1위는 {_top_factor_text(occupation_df, str(industry_meta.get('unit', '')), False)}입니다.",
    ]
    if gy_meta_summary.get("ok"):
        share_now = gy_meta_summary.get("latest_share_pct")
        contrib_now = gy_meta_summary.get("latest_contrib_pct")
        contrib_prev = gy_meta_summary.get("prev_year_contrib_pct")
        contrib_diff = gy_meta_summary.get("contrib_yoy_change_pp")
        share_text = "-" if pd.isna(share_now) else f"{float(share_now):,.2f}%"
        contrib_now_text = "-" if pd.isna(contrib_now) else f"{float(contrib_now):,.1f}%"
        contrib_prev_text = "-" if pd.isna(contrib_prev) else f"{float(contrib_prev):,.1f}%"
        contrib_diff_text = "-" if pd.isna(contrib_diff) else f"{float(contrib_diff):+,.1f}%p"
        summary_lines.append(
            f"전국 취업자 중 {region} 비중은 {share_text}이며, 증감 기여율은 {contrib_prev_text}에서 {contrib_now_text}로 {contrib_diff_text} 변화했습니다."
        )
    top_up, top_down = _pick_top_streak_items()
    if top_up and top_down:
        summary_lines.append(
            f"연속흐름: {_fmt_streak_item(top_up)}, {_fmt_streak_item(top_down)} ({yoy_text}대비 증감 기준)"
        )
    elif top_up:
        summary_lines.append(f"연속흐름: {_fmt_streak_item(top_up)} ({yoy_text}대비 증감 기준)")
    elif top_down:
        summary_lines.append(f"연속흐름: {_fmt_streak_item(top_down)} ({yoy_text}대비 증감 기준)")

    summary_lines.append(_build_new_focus_line("최고", "이번 달 NEW 핵심"))
    summary_lines.append(_build_new_focus_line("최저", "이번 달 NEW 리스크"))

    st.markdown("\n".join([f"- {line}" for line in summary_lines]))

    _report_heading("경제활동인구 현황요약")
    activity_view = activity_df.copy()
    activity_view = activity_view[["지표", "prev_value", "latest_value", "delta_value", "unit"]].rename(
        columns={"prev_value": prev_text, "latest_value": latest_text, "delta_value": f"{yoy_text} 대비 증감"}
    )
    activity_view[prev_text] = activity_view.apply(lambda r: fmt_num(r[prev_text], str(r["unit"])), axis=1)
    activity_view[latest_text] = activity_view.apply(lambda r: fmt_num(r[latest_text], str(r["unit"])), axis=1)
    delta_col = f"{yoy_text} 대비 증감"
    activity_view[delta_col] = activity_view.apply(lambda r: fmt_triangle_delta(r[delta_col], str(r["unit"]), fmt_num), axis=1)
    activity_view = activity_view.drop(columns=["unit"])
    st.dataframe(activity_view, use_container_width=True, hide_index=True)

    _report_heading("취업자수 상세현황")
    structure_lines: List[str] = []
    sections = [("산업별 취업자수", "industry"), ("직종별 취업자수", "occupation"), ("종사상지위별 취업자", "status")]
    for title, ds_key in sections:
        tbl, meta = compute_contribution_table(report_df, region=region, dataset_key=ds_key, lag=lag)
        if not meta.get("ok"):
            st.markdown(f"- **{title}**: 데이터가 부족해 요약을 생성하지 못했습니다.")
            continue
        unit = str(meta.get("unit", ""))
        factor_tbl = _filter_industry_factor_table(tbl) if ds_key == "industry" else tbl
        pos_text = fmt_contrib_items(factor_tbl, unit, positive=True, top_n=3)
        neg_text = fmt_contrib_items(factor_tbl, unit, positive=False, top_n=3)
        line_pos = f"증가요인: {pos_text}"
        line_neg = f"감소요인: {neg_text}"
        streak_line = _build_dataset_streak_summary_line(
            source_df=report_df,
            region=region,
            dataset_key=ds_key,
            asof_period=pd.Timestamp(selected_period),
            indicator_name=str(meta.get("indicator", "")),
            min_len=3,
            yoy_label=yoy_text,
        )
        structure_lines.extend([title, line_pos, line_neg])
        if streak_line:
            structure_lines.append(streak_line)
        st.markdown(f"- **{title}**")
        st.markdown(f"  - {line_pos}")
        st.markdown(f"  - {line_neg}")
        if streak_line:
            st.markdown(f"  - {streak_line}")
        if report_events.empty:
            line_new = f"이번 {labels['point']} NEW 달성: 없음"
            structure_lines.append(line_new)
            st.markdown(f"  - {line_new}")
        else:
            ds_events = report_events[report_events["데이터셋"].astype(str) == str(title)].copy()
            if ds_events.empty:
                line_new = f"이번 {labels['point']} NEW 달성: 없음"
                structure_lines.append(line_new)
                st.markdown(f"  - {line_new}")
            else:
                ds_events = ds_events.sort_values(["구분", "범위", "유형", "분류"], ascending=[True, True, True, True])
                event_tokens: List[str] = []
                for _, er in ds_events.iterrows():
                    cat = str(er.get("분류", "")).strip() or "전체"
                    token = (
                        f"{cat}"
                        f"({str(er.get('구분', ''))} "
                        f"{str(er.get('범위', ''))} "
                        f"{str(er.get('유형', ''))} NEW)"
                    )
                    event_tokens.append(token)
                line_new = f"이번 {labels['point']} NEW 달성: " + ", ".join(event_tokens)
                structure_lines.append(line_new)
                st.markdown(f"  - {line_new}")

    anomaly_df = compute_anomaly_table(report_df, region=region, lag=lag, lookback_periods=36)
    if not anomaly_df.empty:
        anomaly_df = anomaly_df[anomaly_df["기준시점"] == latest_text].copy()
    if not for_pdf:
        st.markdown("##### AI 이상탐지 요약")
        if anomaly_df.empty:
            st.markdown("- 이상탐지 결과가 없습니다.")
        else:
            scores = pd.to_numeric(anomaly_df["이상점수"], errors="coerce")
            focus = anomaly_df[scores >= 50].copy().sort_values("이상점수", ascending=False)
            high_cnt = int((scores >= 75).sum())
            med_cnt = int(((scores >= 50) & (scores < 75)).sum())
            st.markdown(f"- 우선점검(75점 이상): **{high_cnt}건**")
            st.markdown(f"- 주의관찰(50-74점): **{med_cnt}건**")
            if not focus.empty:
                top3 = focus.head(3)
                lines = []
                for _, r in top3.iterrows():
                    lines.append(
                        f"  - [{r['기준시점']}] {r['데이터셋']} / {r['분류']}: "
                        f"{r['이유']} ({float(r['이상점수']):.1f}점)"
                    )
                st.markdown("- Top 3 이벤트\n" + "\n".join(lines))

    _report_heading(f"[참고] 전국대비 {region} 현황")
    reference_lines: List[str] = []
    if region != "전국":
        gy_meta = gy_meta_summary
        if gy_meta.get("ok"):
            ref_prd_se = str(gy_meta.get("prd_se", "M"))
            ref_latest = fmt_period(gy_meta.get("latest_period"), ref_prd_se)
            ref_prev = fmt_period(gy_meta.get("prev_year_period"), ref_prd_se)
            ref_unit = str(gy_meta.get("unit", ""))

            share_now = gy_meta.get("latest_share_pct")
            share_prev = gy_meta.get("prev_year_share_pct")
            share_change = gy_meta.get("share_yoy_change_pp")
            share_now_text = "-" if pd.isna(share_now) else f"{float(share_now):,.2f}%"
            share_prev_text = "-" if pd.isna(share_prev) else f"{float(share_prev):,.2f}%"
            share_change_text = "-" if pd.isna(share_change) else f"{float(share_change):+,.2f}%p"

            contrib_now = gy_meta.get("latest_contrib_pct")
            contrib_prev = gy_meta.get("prev_year_contrib_pct")
            contrib_change = gy_meta.get("contrib_yoy_change_pp")
            contrib_now_text = "-" if pd.isna(contrib_now) else f"{float(contrib_now):,.1f}%"
            contrib_prev_text = "-" if pd.isna(contrib_prev) else f"{float(contrib_prev):,.1f}%"
            contrib_change_text = "-" if pd.isna(contrib_change) else f"{float(contrib_change):+,.1f}%p"

            nat_yoy = gy_meta.get("latest_nat_yoy_abs")
            gg_yoy = gy_meta.get("latest_gg_yoy_abs")
            nat_flow = "증가분" if pd.notna(nat_yoy) and float(nat_yoy) > 0 else ("감소분" if pd.notna(nat_yoy) and float(nat_yoy) < 0 else "변동분")

            reference_lines.append(f"기준시점은 {ref_latest}, 비교시점은 {ref_prev}입니다.")
            reference_lines.append(
                f"{ref_latest} 기준 전국 취업자는 {fmt_num(gy_meta.get('latest_nat_value'), ref_unit)}, "
                f"{region} 취업자는 {fmt_num(gy_meta.get('latest_gg_value'), ref_unit)}입니다."
            )
            reference_lines.append(
                f"전국 대비 {region} 비중은 {ref_prev} {share_prev_text}에서 "
                f"{ref_latest} {share_now_text}로 {share_change_text} 변화했습니다."
            )
            reference_lines.append(
                f"전국 취업자 {nat_flow} 대비 {region} 기여율은 {ref_prev} {contrib_prev_text}에서 "
                f"{ref_latest} {contrib_now_text}로 {contrib_change_text} 변화했습니다."
            )
            reference_lines.append(
                f"같은 시점 전년비 증감은 전국 {fmt_num(nat_yoy, ref_unit)}, "
                f"{region} {fmt_num(gg_yoy, ref_unit)}입니다."
            )
            for line in reference_lines:
                st.markdown(f"- {line}")
            st.caption("참고: 전국 증감이 작을수록 기여율(%)은 크게 흔들릴 수 있습니다.")
        else:
            reference_lines.append("전국 대비 참고 지표를 계산할 수 없습니다.")
            st.markdown("- 전국 대비 참고 지표를 계산할 수 없습니다.")
    else:
        reference_lines.append("전국 선택 시에는 전국대비 지표를 표시하지 않습니다.")
        st.markdown("- 전국 선택 시에는 전국대비 지표를 표시하지 않습니다.")

    st.markdown("---")
    _report_heading("[참고] 취업자수 상세내용")
    detail_sections: List[Dict[str, Any]] = []
    for title, ds_key in sections:
        tbl, meta = compute_contribution_table(report_df, region=region, dataset_key=ds_key, lag=lag)
        st.markdown(f"- **{title}(천명)**")
        if not meta.get("ok") or tbl.empty:
            st.markdown("- 데이터가 부족해 상세 분석을 생성하지 못했습니다.")
            detail_sections.append({"title": f"{title}(천명)", "lines": ["데이터가 부족해 상세 분석을 생성하지 못했습니다."], "table": pd.DataFrame()})
            continue
        unit = str(meta.get("unit", ""))
        latest = fmt_period(meta.get("latest_period"), prd_se)
        prev = fmt_period(meta.get("prev_period"), prd_se)
        total_delta = fmt_num(meta.get("total_delta"), unit)
        factor_tbl = _filter_industry_factor_table(tbl) if ds_key == "industry" else tbl
        top_pos = factor_tbl[factor_tbl["증감"] > 0].nlargest(1, "증감")
        top_neg = factor_tbl[factor_tbl["증감"] < 0].nsmallest(1, "증감")
        pos_line = f"{top_pos.iloc[0]['분류']}({fmt_num(top_pos.iloc[0]['증감'], unit)})" if not top_pos.empty else "없음"
        neg_line = f"{top_neg.iloc[0]['분류']}({fmt_num(top_neg.iloc[0]['증감'], unit)})" if not top_neg.empty else "없음"
        st.markdown(f"- {latest} 기준 총증감: **{total_delta}** ({prev} 대비)")
        st.markdown(f"- 증가요인 1위: **{pos_line}**, 감소요인 1위: **{neg_line}**")
        detail_view = tbl.copy()
        ordered_categories = order_categories_like_ui(
            detail_view["분류"].dropna().astype(str).tolist(),
            ds_key,
            is_gyeonggi31_mode=is_gyeonggi31_mode,
        )
        order_map = {name: idx for idx, name in enumerate(ordered_categories)}
        detail_view["정렬순서"] = detail_view["분류"].map(order_map).fillna(999)
        detail_view = detail_view.sort_values(["정렬순서", "분류"]).drop(columns=["정렬순서"])
        detail_view = detail_view.rename(columns={"최신값": latest, "비교값": prev})
        detail_view["증감"] = detail_view["증감"].apply(lambda x: fmt_triangle_delta(x, unit, fmt_num))
        st.dataframe(detail_view, use_container_width=True, hide_index=True)
        detail_sections.append(
            {
                "title": f"{title}(천명)",
                "lines": [f"{latest} 기준 총증감: {total_delta} ({prev} 대비)", f"증가요인 1위: {pos_line}, 감소요인 1위: {neg_line}"],
                "table": detail_view,
            }
        )

    if not for_pdf:
        st.markdown("##### 점검 액션")
        action_lines = []
        if not industry_factor_df.empty:
            pos = industry_factor_df[industry_factor_df["증감"] > 0].nlargest(1, "증감")
            neg = industry_factor_df[industry_factor_df["증감"] < 0].nsmallest(1, "증감")
            if not pos.empty:
                action_lines.append(f"- 산업 증가 1위({pos.iloc[0]['분류']})의 증가 지속 여부를 다음 {labels['point']}에 점검")
            if not neg.empty:
                action_lines.append(f"- 산업 감소 1위({neg.iloc[0]['분류']})의 구조적 감소 여부를 원인 점검")
        if not anomaly_df.empty:
            high_focus = anomaly_df[pd.to_numeric(anomaly_df["이상점수"], errors="coerce") >= 75]
            action_lines.append(f"- 이상점수 75점 이상 항목 {len(high_focus):,}건에 대해 우선 확인 코멘트 수집")
        if not action_lines:
            action_lines.append("- 이번 시점은 급격한 이상 신호가 제한적이므로 주요 분류 추세 모니터링 유지")
        st.markdown("\n".join(action_lines))

    doc_file_safe_region = re.sub(r"[^0-9A-Za-z가-힣_\\-]", "_", str(region))
    doc_file_safe_period = re.sub(r"[^0-9A-Za-z가-힣_\\-]", "_", str(latest_text))

    printable_html = _build_printable_report_html(
        title=report_title,
        summary_lines=summary_lines,
        activity_table=activity_view,
        structure_lines=structure_lines,
        reference_title=f"[참고] 전국대비 {region} 현황",
        reference_lines=reference_lines,
        detail_sections=detail_sections,
    )
    st.markdown("---")
    st.markdown("#### 인쇄용 HTML 리포트")
    st.caption("아래 영역에서 인쇄 버튼 또는 Ctrl+P로 PDF 저장이 가능합니다.")
    components.html(printable_html, height=1200, scrolling=True)

    doc_bytes = _build_report_docx_bytes(
        title=report_title,
        summary_lines=summary_lines,
        activity_table=activity_view,
        structure_lines=structure_lines,
        reference_title=f"[참고] 전국대비 {region} 현황",
        reference_lines=reference_lines,
        detail_sections=detail_sections,
    )
    st.download_button(
        "DOCX 다운로드",
        data=doc_bytes,
        file_name=f"{doc_file_safe_period}_{doc_file_safe_region}_경제활동인구_브리프.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        key="download_report_docx",
    )
